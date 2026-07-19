"""Aşama 1 — Belge Okuma (Parsing).

Belge Hazırlama Sürecinin (Ingestion Pipeline) ilk adımı: ham dosyaları
programın işleyebileceği düz metne çevirmek.

Bu naive sürümde desteklenenler:
- .txt / .md  → düz metin, tek kayıt
- .pdf        → pypdfium2 ile sayfa sayfa (sayfa numarası sitasyon için saklanır)
- .docx       → python-docx ile paragraf + tablo
- .xlsx       → openpyxl ile sayfa (sheet) sayfa; tablo satırları metne çevrilir

OCR yok: taranmış (görüntü) PDF'lerin metin katmanı boş döner.
PowerPoint ve eski .doc/.xls formatları yok (bkz. gelistirme_odevleri.md).

Çalıştır:  python asama1_belge_okuma.py
"""

from pathlib import Path

DOKUMAN_KLASORU = Path(__file__).parent / "dokumanlar"


def metin_dosyasi_oku(yol: Path) -> list[dict]:
    """Düz metin dosyasını tek kayıt olarak döndürür."""
    metin = yol.read_text(encoding="utf-8")
    return [{"dosya": yol.name, "sayfa": None, "metin": metin}]


def pdf_oku(yol: Path) -> list[dict]:
    """PDF'i sayfa sayfa okur.

    Sayfa numarasını atmıyoruz: yanıtta "hangi dosyanın hangi sayfası"
    diyebilmek (citation) bu bilgiye bağlı.
    """
    import pypdfium2 as pdfium

    kayitlar = []
    belge = pdfium.PdfDocument(yol)
    for sayfa_no, sayfa in enumerate(belge, start=1):
        metin = sayfa.get_textpage().get_text_range()
        kayitlar.append({"dosya": yol.name, "sayfa": sayfa_no, "metin": metin})
    belge.close()
    return kayitlar


def word_oku(yol: Path) -> list[dict]:
    """Word belgesini okur: paragraflar + tablolar, tek kayıt.

    Tablo hücreleri " | " ile birleştirilir — naive ama işlevsel;
    tablo yapısını koruyan zengin ayrıştırma geliştirme ödevidir.
    """
    import docx

    belge = docx.Document(yol)
    satirlar = [p.text for p in belge.paragraphs if p.text.strip()]
    for tablo in belge.tables:
        for satir in tablo.rows:
            satirlar.append(" | ".join(h.text.strip() for h in satir.cells))
    return [{"dosya": yol.name, "sayfa": None, "metin": "\n".join(satirlar)}]


def excel_oku(yol: Path) -> list[dict]:
    """Excel dosyasını sayfa (sheet) sayfa okur; her satır " | " ile metne çevrilir.

    Sheet adı "sayfa" metadata'sına yazılır → sitasyon [dosya, sayfa Otoyollar]
    biçiminde çıkar. Başlık satırı metnin başında kalır ki parça tek başına
    okunduğunda kolonların anlamı kaybolmasın.
    """
    import openpyxl

    kayitlar = []
    kitap = openpyxl.load_workbook(yol, data_only=True)
    for sayfa in kitap.worksheets:
        satirlar = []
        for satir in sayfa.iter_rows(values_only=True):
            hucreler = [str(h) for h in satir if h is not None]
            if hucreler:
                satirlar.append(" | ".join(hucreler))
        if satirlar:
            kayitlar.append({"dosya": yol.name, "sayfa": sayfa.title,
                             "metin": "\n".join(satirlar)})
    return kayitlar


def belgeleri_oku(klasor: Path = DOKUMAN_KLASORU) -> list[dict]:
    """Klasördeki tüm belgeleri okur.

    Dönen her kayıt: {"dosya": ..., "sayfa": int | str | None, "metin": ...}
    ("sayfa": PDF'te sayfa numarası, Excel'de sheet adı, diğerlerinde None)

    "dosya" ve "sayfa" alanları METADATA'dır (belge metadata'sı):
    metnin kendisi değil, metin HAKKINDA bilgi. Doğru kaynak göstermenin
    temeli daha bu ilk aşamada atılır.
    """
    okuyucular = {
        ".txt": metin_dosyasi_oku,
        ".md": metin_dosyasi_oku,
        ".pdf": pdf_oku,
        ".docx": word_oku,
        ".xlsx": excel_oku,
    }
    kayitlar = []
    for yol in sorted(klasor.iterdir()):
        okuyucu = okuyucular.get(yol.suffix.lower())
        if okuyucu:
            kayitlar += okuyucu(yol)
    return kayitlar


if __name__ == "__main__":
    kayitlar = belgeleri_oku()
    print(f"{len(kayitlar)} kayıt okundu.\n")
    for k in kayitlar:
        sayfa = f"sayfa {k['sayfa']}" if k["sayfa"] else "tek parça"
        print(f"- {k['dosya']:<32} {sayfa:<10} {len(k['metin']):>6} karakter")
